import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.document import Document
from app.models.subject import Subject
from app.models.user import User
from app.routers.auth import get_current_user
from app.schemas.document_schema import (
    DocumentCreate,
    DocumentResponse,
    DocumentUpdate,
)


router = APIRouter(
    prefix="/api/documents",
    tags=["Documents"],
)

UPLOAD_ROOT = Path("uploads/documents")
MAX_UPLOAD_SIZE = 10 * 1024 * 1024

ALLOWED_FILE_EXTENSIONS = {
    ".txt",
    ".pdf",
    ".docx",
}


def safe_filename(filename: str) -> str:
    filename = filename.strip() or "document"
    filename = re.sub(r"[^\w.\-À-ỹ ]", "_", filename)
    filename = filename.replace(" ", "_")
    return filename


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))

    pages_text: list[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text.strip())

    return "\n\n".join(pages_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = WordDocument(BytesIO(file_bytes))

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    return "\n\n".join(paragraphs)


def extract_text_from_uploaded_file(file_bytes: bytes, extension: str) -> str:
    if extension == ".txt":
        return extract_text_from_txt(file_bytes)

    if extension == ".pdf":
        return extract_text_from_pdf(file_bytes)

    if extension == ".docx":
        return extract_text_from_docx(file_bytes)

    return ""


def save_uploaded_file(
    user_id: int,
    original_filename: str,
    file_bytes: bytes,
) -> str:
    user_upload_dir = UPLOAD_ROOT / str(user_id)
    user_upload_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
    stored_filename = f"{timestamp}_{safe_filename(original_filename)}"

    file_path = user_upload_dir / stored_filename
    file_path.write_bytes(file_bytes)

    return str(file_path)


def clean_optional_text(value: str | None) -> str | None:
    if value is None:
        return None

    cleaned_value = value.strip()
    return cleaned_value if cleaned_value else None


def count_words(content: str) -> int:
    if not content or not content.strip():
        return 0

    return len(content.strip().split())


def build_chunks(content: str, max_length: int = 450) -> list[str]:
    if not content or not content.strip():
        return []

    clean_content = re.sub(r"\s+", " ", content).strip()

    sentences = re.split(r"(?<=[.!?。！？])\s+", clean_content)
    sentences = [sentence.strip() for sentence in sentences if sentence.strip()]

    chunks: list[str] = []
    current_chunk = ""

    for sentence in sentences:
        candidate = f"{current_chunk} {sentence}".strip() if current_chunk else sentence

        if len(candidate) > max_length and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = sentence
        else:
            current_chunk = candidate

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks if chunks else [clean_content]


def get_subject_for_current_user(
    db: Session,
    subject_id: int,
    user_id: int,
) -> Subject:
    subject = db.scalar(
        select(Subject).where(
            Subject.id == subject_id,
            Subject.user_id == user_id,
        )
    )

    if not subject:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Không tìm thấy môn học của bạn.",
        )

    return subject


def get_document_for_current_user(
    db: Session,
    document_id: int,
    user_id: int,
) -> tuple[Document, Subject]:
    result = db.execute(
        select(Document, Subject)
        .join(Subject, Document.subject_id == Subject.id)
        .where(
            Document.id == document_id,
            Subject.user_id == user_id,
        )
    ).first()

    if not result:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Không tìm thấy tài liệu.",
        )

    document, subject = result
    return document, subject


def document_to_response(
    document: Document,
    subject: Subject,
) -> DocumentResponse:
    chunks = build_chunks(document.content)

    return DocumentResponse(
        id=document.id,
        subject_id=document.subject_id,
        subject_name=subject.name,
        title=document.title,
        source_type=document.source_type,
        file_name=document.file_name,
        file_path=document.file_path,
        source_url=document.source_url,
        content=document.content,
        summary=document.summary,
        keywords=document.keywords,
        relevance_score=document.relevance_score,
        processing_status=document.processing_status,
        chunks=chunks,
        word_count=count_words(document.content),
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


@router.get("", response_model=list[DocumentResponse])
def get_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    rows = db.execute(
        select(Document, Subject)
        .join(Subject, Document.subject_id == Subject.id)
        .where(Subject.user_id == current_user.id)
        .order_by(Document.created_at.desc())
    ).all()

    return [
        document_to_response(document, subject)
        for document, subject in rows
    ]


@router.post(
    "",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_document(
    data: DocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    subject = get_subject_for_current_user(
        db=db,
        subject_id=data.subject_id,
        user_id=current_user.id,
    )

    document = Document(
        subject_id=subject.id,
        title=data.title.strip(),
        source_type=data.source_type.strip().upper(),
        file_name=clean_optional_text(data.file_name),
        file_path=clean_optional_text(data.file_path),
        source_url=clean_optional_text(data.source_url),
        content=data.content.strip(),
        summary=clean_optional_text(data.summary),
        keywords=clean_optional_text(data.keywords),
        relevance_score=data.relevance_score,
        processing_status="DA_XU_LY",
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    return document_to_response(document, subject)

@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document_file(
    subject_id: int = Form(...),
    title: str | None = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    subject = get_subject_for_current_user(
        db=db,
        subject_id=subject_id,
        user_id=current_user.id,
    )

    original_filename = file.filename or "document"
    extension = Path(original_filename).suffix.lower()

    if extension == ".doc":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File .doc cũ chưa được hỗ trợ. Vui lòng chuyển sang .docx rồi tải lại.",
        )

    if extension not in ALLOWED_FILE_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chỉ hỗ trợ file .txt, .pdf hoặc .docx.",
        )

    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File tải lên đang trống.",
        )

    if len(file_bytes) > MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File quá lớn. Vui lòng tải file nhỏ hơn 10MB.",
        )

    try:
        content = extract_text_from_uploaded_file(
            file_bytes=file_bytes,
            extension=extension,
        )
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Không thể đọc nội dung file. Vui lòng kiểm tra lại file.",
        )

    if not content.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Không đọc được nội dung trong file. Nếu là PDF scan ảnh thì hệ thống chưa hỗ trợ OCR.",
        )

    saved_file_path = save_uploaded_file(
        user_id=current_user.id,
        original_filename=original_filename,
        file_bytes=file_bytes,
    )

    document_title = title.strip() if title and title.strip() else Path(original_filename).stem

    document = Document(
        subject_id=subject.id,
        title=document_title,
        source_type="FILE",
        file_name=original_filename,
        file_path=saved_file_path,
        source_url=None,
        content=content.strip(),
        summary=None,
        keywords=None,
        relevance_score=None,
        processing_status="DA_XU_LY",
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    return document_to_response(document, subject)


@router.put("/{document_id}", response_model=DocumentResponse)
def update_document(
    document_id: int,
    data: DocumentUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document, _old_subject = get_document_for_current_user(
        db=db,
        document_id=document_id,
        user_id=current_user.id,
    )

    subject = get_subject_for_current_user(
        db=db,
        subject_id=data.subject_id,
        user_id=current_user.id,
    )

    document.subject_id = subject.id
    document.title = data.title.strip()
    document.source_type = data.source_type.strip().upper()
    document.file_name = clean_optional_text(data.file_name)
    document.file_path = clean_optional_text(data.file_path)
    document.source_url = clean_optional_text(data.source_url)
    document.content = data.content.strip()
    document.summary = clean_optional_text(data.summary)
    document.keywords = clean_optional_text(data.keywords)
    document.relevance_score = data.relevance_score
    document.processing_status = "DA_XU_LY"

    db.commit()
    db.refresh(document)

    return document_to_response(document, subject)


@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document, _subject = get_document_for_current_user(
        db=db,
        document_id=document_id,
        user_id=current_user.id,
    )

    db.delete(document)
    db.commit()

    return {
        "message": "Xóa tài liệu thành công.",
    }